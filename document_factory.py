from abc import ABC, abstractmethod
import PyPDF2
from docx import Document as DocxDocument

class Document(ABC):
    @abstractmethod
    def read_document(self, filename):
        pass

    @abstractmethod
    def write_document(self, filename):
        pass

class PDF(Document):
    def read_document(self, filename):
        pdf = PyPDF2.PdfReader(open(filename, 'rb'))
        print('PDF document read')

    def write_document(self, filename):
        # You need to provide a PdfFileWriter object to write to a PDF file
        pdf_writer = PyPDF2.PdfWriter()
        with open(filename, 'wb') as file:
            pdf_writer.write(file)
        print('PDF document written')

class Word(Document):
    def read_document(self, filename):
        document = DocxDocument(filename)
        for paragraph in document.paragraphs:
            print(paragraph.text)

    def write_document(self, filename):
        document = DocxDocument()
        document.add_paragraph('Hello, world!')
        document.save(filename)
        print('Word document written')

class DocumentFactory(ABC):
    @abstractmethod
    def create_document(self):
        pass


class PDFFactory(DocumentFactory):
    def create_document(self):
        return PDF()


class WordFactory(DocumentFactory):
    def create_document(self):
        return Word()


class DocumentFactoryProducer:
    def __init__(self, factory):
        self.__factory = factory

    def create_document(self):
        return self.__factory.create_document()


if __name__ == '__main__':
    # Create a PDF document
    pdf_factory = PDFFactory()
    pdf_document = pdf_factory.create_document()
    pdf_document.write_document('test.pdf')
    pdf_document.read_document('test.pdf')

    # Create a Word document
    word_factory = WordFactory()
    word_document = word_factory.create_document()
    word_document.write_document('test.docx')
    word_document.read_document('test.docx')
